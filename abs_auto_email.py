import smtplib
import os
import sys
import subprocess
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document
from datetime import datetime, timedelta

#이메일 설정
EMAIL_ADDRESS = os.environ.get("EMAIL_ADDRESS")
EMAIL_PASSWORD = os.environ.get("EMAIL_PASSWORD")
RECEIVER_EMAIL = "trotz4210@gmail.com" #총동연: "khudongari@khu.ac.kr"
TEMPLATE_PATH = "template.docx"

TODAY = datetime.now()
TODAY = TODAY.strftime("%Y년  %m월  %d일")
NOW = datetime.now() + timedelta(days=7)
NOW = NOW.strftime("%Y-%m-%d")

#기준일(월요일)부터 7일간 일시 텍스트 생성
def gen_dates():
    today = datetime.now()
    weekdays_li = ["월요일", "화요일", "수요일", "목요일", "금요일", "토요일", "일요일"]
    dates = []
    for i in range(1, 8):
        standard_date = today + timedelta(days=i+3)
        date_str = standard_date.strftime("%Y년 %m월 %d일")
        weekdays_str = weekdays_li[standard_date.weekday()]

        full_date = f"{date_str} {weekdays_str} 23:00-07:00"
        dates.append(full_date)
    return dates

def get_Nth_week(date_obj):
    """해당 월의 '첫 번째 월요일'이 포함된 주를 1주차로 계산"""
    # 에러 방지: 만약 문자열(str)이 들어오면 datetime 객체로 강제 변환
    if isinstance(date_obj, str):
        # YYYY-MM-DD 형식의 문자열일 경우 변환
        date_obj = datetime.strptime(date_obj[:10], "%Y-%m-%d")

    # 1. 해당 월의 1일 구하기 (이제 에러 발생 안 함)
    first_day = date_obj.replace(day=1)
    
    # 2. 첫 번째 월요일 찾기
    dist_to_monday = (0 - first_day.weekday() + 7) % 7
    first_monday = first_day + timedelta(days=dist_to_monday)
    
    # 3. 주차 계산
    diff_days = (date_obj - first_monday).days
    
    if diff_days < 0:
        last_day_of_prev_month = first_day - timedelta(days=1)
        return get_Nth_week(last_day_of_prev_month)
    
    month_week = (diff_days // 7) + 1
    return f"{date_obj.month}월 {month_week}주"

#워드 파일에 표 삽입
def modify_word_document():

    if not os.path.exists(TEMPLATE_PATH):
        raise FileExistsError(f"파일을 찾을 수 없습니다.: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)
    week_info = get_Nth_week(NOW)
    
    #문서 내 중반부 날짜 변경
    target_keyword = "INSERT DATE"

    for paragraph in doc.paragraphs:
        if target_keyword in paragraph.text:
            for run in paragraph.runs:
                if target_keyword in run.text:
                    run.text = run.text.replace(run.text.strip(), TODAY)

    #문서 내 첫째 줄 제목 변경
    title_para = doc.paragraphs[0]
    if title_para.runs:
        new_title = f"{week_info} 야간 동아리방 이용 사전신청서"

        title_para.runs[0].text = new_title

        for i in range(1, len(title_para.runs)):
            title_para.runs[i].text = ""

    #문서 내 3번째 표 선택
    dates = gen_dates()
    table = doc.tables[2]

    TARGET_COL = 0

    for i, date_text in enumerate(dates):
        row_index = i + 2

        if row_index < len(table.rows):
            target_cell = table.rows[row_index].cells[TARGET_COL]
            target_cell.text = date_text
            target_cell.paragraphs[0].text = date_text #폰트, 크기 유지

    output_filename = f"[ABS] {get_Nth_week(NOW)} 야간 동아리방 이용 사전신청서.docx"
    doc.save(output_filename)
    return output_filename

def convert_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    print(f"Start to convert: {docx_path} -> {pdf_path}")

    if sys.platform == "win32":
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        except Exception as e:
            print("Failed to convert Windows PDF")
            raise e
    else:
        try:
            subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", docx_path], check=True)
        except Exception as e:
            print("Failed to convert Linux PDF")
            raise e
    return pdf_path

def send_email(attachment_path):

    msg = MIMEMultipart()
    msg['From'] = EMAIL_ADDRESS
    msg['To'] = RECEIVER_EMAIL
    msg['Subject'] = f"[ABS] {get_Nth_week(NOW)} 야간 동아리방 이용 사전신청서"

    body = f"안녕하십니까. 주식경제동아리 ABS입니다. \n\n{get_Nth_week(NOW)} 야간 동아리방 이용 사전신청서 송부드립니다."
    msg.attach(MIMEText(body, 'plain'))

    with open(attachment_path, "rb") as f:
        part = MIMEApplication(f.read(), Name = os.path.basename(attachment_path))
        part['Content-Disposition'] = f'attachment; filename="{os.path.basename(attachment_path)}"'
        msg.attach(part)

    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
    server.send_message(msg)
    server.quit()

# Running--
if __name__ == "__main__":
    try:
        # word파일 생성
        new_doc_path = modify_word_document()
        # pdf 변환
        pdf_file_path = convert_to_pdf(new_doc_path)
        send_email(pdf_file_path)
        print("Success")

    except Exception as e:
        print(f"Error: {e}")
